import sys
import os

sys.path.append(os.path.join(os.path.dirname(__file__), '..', '..'))

try:
    from ..models.job_analyzer import JobAnalyzer
    from ..models.profile_manager import ProfileManager
    from ..models.resume_generator import ResumeGenerator
    from ..database.db_manager import DatabaseManager
except ImportError as e:
    print(f"Error importing Job Seeker modules: {e}")
    print(f"Current working directory: {os.getcwd()}")
    print(f"Python path: {sys.path}")
    raise

class JobSeekerIntegration:
    def __init__(self):
        # Set correct data directory path for job seeker models
        data_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'data')
        
        self.job_analyzer = JobAnalyzer()
        self.profile_manager = ProfileManager(data_dir)
        self.resume_generator = ResumeGenerator()
        self.db_manager = DatabaseManager()
        self.current_analysis = None
        self.current_profile = None

    def analyze_job(self, job_description):
        try:
            if not job_description.strip():
                return {
                    'success': False,
                    'message': 'Job description is required',
                    'data': None,
                    'errors': ['Empty job description']
                }

            keywords = self.job_analyzer.extract_keywords(job_description)
            requirements = self.job_analyzer.identify_requirements(job_description)

            analysis = {
                'keywords': keywords,
                'relevance_score': 0.0,
                **requirements
            }

            self.current_analysis = analysis

            # Database saving is optional for job analysis
            # Removed to keep the system simple and focused

            return {
                'success': True,
                'message': 'Job analysis completed successfully',
                'data': analysis,
                'errors': []
            }
        except Exception as e:
            return {
                'success': False,
                'message': 'Error analyzing job description',
                'data': None,
                'errors': [str(e)]
            }
    
    def get_profile(self):
        """Get current profile data"""
        try:
            profile = self.profile_manager.load_profile()
            
            if profile:
                self.current_profile = profile
                return {
                    'success': True,
                    'message': 'Profile loaded successfully',
                    'data': profile,
                    'errors': []
                }
            else:
                return {
                    'success': True,
                    'message': 'No profile found',
                    'data': None,
                    'errors': []
                }
        except Exception as e:
            return {
                'success': False,
                'message': 'Error loading profile',
                'data': None,
                'errors': [str(e)]
            }
    
    def create_profile(self, profile_data):
        """Create a new profile"""
        try:
            # Validate profile data
            errors = self.profile_manager.validate_profile(profile_data)
            if errors:
                return {
                    'success': False,
                    'message': 'Profile validation failed',
                    'data': None,
                    'errors': errors
                }
            
            success = self.profile_manager.create_profile(profile_data)
            if success:
                self.current_profile = profile_data
                return {
                    'success': True,
                    'message': 'Profile created successfully',
                    'data': profile_data,
                    'errors': []
                }
            else:
                return {
                    'success': False,
                    'message': 'Failed to create profile',
                    'data': None,
                    'errors': ['Profile creation failed']
                }
        except Exception as e:
            return {
                'success': False,
                'message': 'Error creating profile',
                'data': None,
                'errors': [str(e)]
            }
    
    def update_profile(self, profile_data):
        """Update existing profile"""
        try:
            # Validate profile data
            errors = self.profile_manager.validate_profile(profile_data)
            if errors:
                return {
                    'success': False,
                    'message': 'Profile validation failed',
                    'data': None,
                    'errors': errors
                }
            
            success = self.profile_manager.update_profile(profile_data)
            if success:
                self.current_profile = profile_data
                return {
                    'success': True,
                    'message': 'Profile updated successfully',
                    'data': profile_data,
                    'errors': []
                }
            else:
                return {
                    'success': False,
                    'message': 'Failed to update profile',
                    'data': None,
                    'errors': ['Profile update failed']
                }
        except Exception as e:
            return {
                'success': False,
                'message': 'Error updating profile',
                'data': None,
                'errors': [str(e)]
            }
    
    def validate_profile(self, profile_data):
        """Validate profile data"""
        try:
            errors = self.profile_manager.validate_profile(profile_data)
            return {
                'success': len(errors) == 0,
                'message': 'Profile is valid' if len(errors) == 0 else 'Profile validation failed',
                'data': {'valid': len(errors) == 0, 'errors': errors},
                'errors': errors
            }
        except Exception as e:
            return {
                'success': False,
                'message': 'Error validating profile',
                'data': None,
                'errors': [str(e)]
            }
    
    def generate_resume(self, profile=None, analysis=None):
        """Generate tailored resume"""
        try:
            if not profile:
                profile = self.current_profile
            if not analysis:
                analysis = self.current_analysis
            
            if not profile:
                return {
                    'success': False,
                    'message': 'Profile is required for resume generation',
                    'data': None,
                    'errors': ['No profile available']
                }
            
            if not analysis:
                return {
                    'success': False,
                    'message': 'Job analysis is required for resume generation',
                    'data': None,
                    'errors': ['No job analysis available']
                }
            
            resume_data = self.resume_generator.generate_resume(profile, analysis)
            formatted_resume = self.resume_generator.format_resume(resume_data)
            
            return {
                'success': True,
                'message': 'Resume generated successfully',
                'data': {
                    'resume_data': resume_data,
                    'formatted_resume': formatted_resume
                },
                'errors': []
            }
        except Exception as e:
            return {
                'success': False,
                'message': 'Error generating resume',
                'data': None,
                'errors': [str(e)]
            }
    
    def export_resume(self, content, format_type, profile=None, job_description=None):
        """Export resume in specified format for web download"""
        try:
            import tempfile
            import os
            from datetime import datetime
            
            if not content:
                return {
                    'success': False,
                    'message': 'No resume content available for export',
                    'data': None,
                    'errors': ['No resume content']
                }
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_name = "resume"
            if profile and "personal_info" in profile:
                name = profile["personal_info"].get("name", "").strip()
                if name:
                    default_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()
                    default_name = default_name.replace(' ', '_').lower()
            
            filename = f"{default_name}_{timestamp}.{format_type}"
            
            # Generate file content based on format
            if format_type == 'txt':
                file_data = content.encode('utf-8')
                content_type = 'text/plain'
            elif format_type == 'pdf':
                file_data = self._generate_pdf(content)
                content_type = 'application/pdf'
            elif format_type == 'docx':
                file_data = self._generate_docx(content)
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                return {
                    'success': False,
                    'message': 'Unsupported export format',
                    'data': None,
                    'errors': ['Format must be txt, pdf, or docx']
                }
            
            return {
                'success': True,
                'message': 'Resume exported successfully',
                'data': {
                    'filename': filename,
                    'file_data': file_data,
                    'content_type': content_type,
                    'format': format_type
                },
                'errors': []
            }
        except Exception as e:
            return {
                'success': False,
                'message': 'Error exporting resume',
                'data': None,
                'errors': [str(e)]
            }
    
    def _generate_pdf(self, content):
        """Generate PDF content"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Check if it's a header (all caps or starts with specific patterns)
                    if para.isupper() or para.startswith('=') or para.startswith('-'):
                        story.append(Paragraph(para.strip(), styles['Heading2']))
                    else:
                        story.append(Paragraph(para.strip(), styles['Normal']))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except ImportError:
            # Fallback to text if reportlab not available
            return content.encode('utf-8')
        except Exception as e:
            print(f"PDF generation error: {e}")
            return content.encode('utf-8')
    
    def _generate_docx(self, content):
        """Generate DOCX content"""
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document()
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                # Check if it's a header
                if para.isupper() or para.startswith('=') or para.startswith('-'):
                    doc.add_heading(para.replace('=', '').replace('-', '').strip(), level=2)
                else:
                    doc.add_paragraph(para)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except ImportError:
            # Fallback to text if python-docx not available
            return content.encode('utf-8')
        except Exception as e:
            print(f"DOCX generation error: {e}")
            return content.encode('utf-8')