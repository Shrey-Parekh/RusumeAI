"""
Export Manager - Handles resume export in multiple formats (PDF, Word, Text)
"""

import os
import json
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class ExportManager:
    """Manages resume export functionality and history tracking"""
    
    def __init__(self, data_dir: str = "data"):
        """Initialize export manager with data directory"""
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(exist_ok=True)
        self.history_file = self.data_dir / "resume_history.json"
        self.resume_history = self._load_history()
    
    def export_to_pdf(self, resume_content: str, filename: str, 
                     profile_data: Optional[Dict[str, Any]] = None,
                     job_description: Optional[str] = None) -> bool:
        """Export resume to PDF format"""
        try:
            if not REPORTLAB_AVAILABLE:
                messagebox.showerror("PDF Export Error", 
                                   "PDF export requires reportlab library.\n"
                                   "Install with: pip install reportlab")
                return False
            
            # Ensure filename has .pdf extension
            if not filename.lower().endswith('.pdf'):
                filename += '.pdf'
            
            # Create PDF document
            doc = SimpleDocTemplate(filename, pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            # Build PDF content
            story = []
            
            # Split content into sections and format
            lines = resume_content.split('\n')
            current_section = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    if current_section:
                        # Process current section
                        section_text = '\n'.join(current_section)
                        if current_section[0].isupper() or current_section[0].startswith('='):
                            # This is likely a header
                            story.append(Paragraph(section_text, styles['Heading2']))
                        else:
                            story.append(Paragraph(section_text, styles['Normal']))
                        story.append(Spacer(1, 12))
                        current_section = []
                else:
                    current_section.append(line)
            
            # Add remaining content
            if current_section:
                section_text = '\n'.join(current_section)
                story.append(Paragraph(section_text, styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            # Add to history
            self._add_to_history(filename, "pdf", profile_data, job_description)
            
            return True
            
        except Exception as e:
            messagebox.showerror("PDF Export Error", f"Failed to export PDF: {str(e)}")
            return False
    
    def export_to_docx(self, resume_content: str, filename: str,
                      profile_data: Optional[Dict[str, Any]] = None,
                      job_description: Optional[str] = None) -> bool:
        """Export resume to Word document format"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                messagebox.showerror("Word Export Error", 
                                   "Word export requires python-docx library.\n"
                                   "Install with: pip install python-docx")
                return False
            
            # Ensure filename has .docx extension
            if not filename.lower().endswith('.docx'):
                filename += '.docx'
            
            # Create Word document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Split content into paragraphs and format
            lines = resume_content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Determine if this is a header (all caps, starts with =, or other indicators)
                if (line.isupper() and len(line) < 50) or line.startswith('=') or line.startswith('---'):
                    # Add as heading
                    heading = doc.add_heading(line.replace('=', '').strip(), level=2)
                else:
                    # Add as normal paragraph
                    doc.add_paragraph(line)
            
            # Save document
            doc.save(filename)
            
            # Add to history
            self._add_to_history(filename, "docx", profile_data, job_description)
            
            return True
            
        except Exception as e:
            messagebox.showerror("Word Export Error", f"Failed to export Word document: {str(e)}")
            return False
    
    def export_to_txt(self, resume_content: str, filename: str,
                     profile_data: Optional[Dict[str, Any]] = None,
                     job_description: Optional[str] = None) -> bool:
        """Export resume to text format"""
        try:
            # Ensure filename has .txt extension
            if not filename.lower().endswith('.txt'):
                filename += '.txt'
            
            # Write content to file
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(resume_content)
            
            # Add to history
            self._add_to_history(filename, "txt", profile_data, job_description)
            
            return True
            
        except Exception as e:
            messagebox.showerror("Text Export Error", f"Failed to export text file: {str(e)}")
            return False
    
    def show_export_dialog(self, resume_content: str, 
                          profile_data: Optional[Dict[str, Any]] = None,
                          job_description: Optional[str] = None) -> Optional[str]:
        """Show file save dialog and export resume"""
        try:
            # Get default filename based on profile name and timestamp
            default_name = "resume"
            if profile_data and "personal_info" in profile_data:
                name = profile_data["personal_info"].get("name", "").strip()
                if name:
                    # Clean name for filename
                    default_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()
                    default_name = default_name.replace(' ', '_').lower()
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"{default_name}_{timestamp}"
            
            # Show file dialog
            filetypes = [
                ("Text files", "*.txt"),
                ("All files", "*.*")
            ]
            
            # Add PDF and Word options if libraries are available
            if REPORTLAB_AVAILABLE:
                filetypes.insert(0, ("PDF files", "*.pdf"))
            if PYTHON_DOCX_AVAILABLE:
                filetypes.insert(-1, ("Word documents", "*.docx"))
            
            filename = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=filetypes,
                initialfile=default_filename,
                title="Export Resume"
            )
            
            if not filename:
                return None  # User cancelled
            
            # Determine export format based on extension
            extension = Path(filename).suffix.lower()
            
            success = False
            if extension == '.pdf':
                success = self.export_to_pdf(resume_content, filename, profile_data, job_description)
            elif extension == '.docx':
                success = self.export_to_docx(resume_content, filename, profile_data, job_description)
            else:
                success = self.export_to_txt(resume_content, filename, profile_data, job_description)
            
            if success:
                messagebox.showinfo("Export Successful", f"Resume exported successfully to:\n{filename}")
                return filename
            else:
                return None
                
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export resume: {str(e)}")
            return None
    
    def get_resume_history(self) -> List[Dict[str, Any]]:
        """Get list of previously exported resumes"""
        return self.resume_history.copy()
    
    def clear_history(self) -> bool:
        """Clear resume export history"""
        try:
            self.resume_history = []
            self._save_history()
            return True
        except Exception as e:
            print(f"Error clearing history: {e}")
            return False
    
    def delete_history_entry(self, entry_id: str) -> bool:
        """Delete a specific history entry"""
        try:
            self.resume_history = [entry for entry in self.resume_history 
                                 if entry.get("id") != entry_id]
            self._save_history()
            return True
        except Exception as e:
            print(f"Error deleting history entry: {e}")
            return False
    
    def _load_history(self) -> List[Dict[str, Any]]:
        """Load resume export history from file"""
        try:
            if self.history_file.exists():
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            return []
        except Exception as e:
            print(f"Error loading history: {e}")
            return []
    
    def _save_history(self) -> None:
        """Save resume export history to file"""
        try:
            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(self.resume_history, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Error saving history: {e}")
    
    def _add_to_history(self, filename: str, format_type: str,
                       profile_data: Optional[Dict[str, Any]] = None,
                       job_description: Optional[str] = None) -> None:
        """Add export to history"""
        try:
            entry = {
                "id": f"{datetime.now().isoformat()}_{len(self.resume_history)}",
                "filename": filename,
                "format": format_type,
                "timestamp": datetime.now().isoformat(),
                "profile_name": "",
                "job_title": "",
                "file_exists": os.path.exists(filename)
            }
            
            # Extract profile name if available
            if profile_data and "personal_info" in profile_data:
                entry["profile_name"] = profile_data["personal_info"].get("name", "")
            
            # Extract job title from job description if available
            if job_description:
                # Simple extraction - look for common job title patterns
                lines = job_description.split('\n')[:5]  # Check first 5 lines
                for line in lines:
                    line = line.strip()
                    if any(keyword in line.lower() for keyword in ['position', 'role', 'job title', 'title:']):
                        entry["job_title"] = line[:100]  # Limit length
                        break
                    elif len(line) < 100 and any(keyword in line.lower() for keyword in 
                                               ['developer', 'engineer', 'manager', 'analyst', 'specialist']):
                        entry["job_title"] = line
                        break
            
            # Add to history (keep last 50 entries)
            self.resume_history.append(entry)
            if len(self.resume_history) > 50:
                self.resume_history = self.resume_history[-50:]
            
            self._save_history()
            
        except Exception as e:
            print(f"Error adding to history: {e}")
    
    @staticmethod
    def check_dependencies() -> Dict[str, bool]:
        """Check which export dependencies are available"""
        return {
            "pdf": REPORTLAB_AVAILABLE,
            "docx": PYTHON_DOCX_AVAILABLE,
            "txt": True  # Always available
        }
    
    @staticmethod
    def get_missing_dependencies() -> List[str]:
        """Get list of missing dependencies for export formats"""
        missing = []
        if not REPORTLAB_AVAILABLE:
            missing.append("reportlab (for PDF export)")
        if not PYTHON_DOCX_AVAILABLE:
            missing.append("python-docx (for Word export)")
        return missing